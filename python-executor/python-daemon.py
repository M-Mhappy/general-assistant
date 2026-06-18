"""
Python 长驻子进程 — JSON-line 协议通信。

协议（每行一个 JSON 对象）：
  接收（stdin）:
    {"type":"execute","code":"...","id":"req-1"}
    {"type":"install","package":"requests","id":"req-2"}
    {"type":"cancel","id":"req-3"}
    {"type":"shutdown"}

  发送（stdout）:
    {"type":"ready"}
    {"type":"stream_output","id":"req-1","stream":"stdout","text":"hello\\n"}
    {"type":"stream_output","id":"req-1","stream":"stderr","text":"..."}
    {"type":"executeResult","id":"req-1","result":"...","files":["..."]}
    {"type":"executeResult","id":"req-1","error":"traceback..."}
    {"type":"error","message":"..."}
    {"type":"shutdownAck"}
"""

import sys
import json
import traceback
import io
import os
import ast
import importlib
import subprocess as sp
from pathlib import Path

# 强制 UTF-8 编码，确保与 Node.js 主进程的 JSON 通信编码一致。
# Windows 上默认可能为 GBK，会导致中文字符乱码或 UnicodeDecodeError。
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

# 保存真实 stdout 引用，用于协议消息发送（不受 StreamProxy 影响）
_REAL_STDOUT = sys.stdout


def send_message(msg: dict):
    """向真实 stdout 发送一条 JSON-line 消息，立即刷新。

    注意：必须使用 _REAL_STDOUT 而非 sys.stdout，
    因为在代码执行期间 sys.stdout 会被替换为 StreamProxy，
    若使用后者会导致 send_message → write → send_message 无限递归。
    """
    line = json.dumps(msg, ensure_ascii=False)
    _REAL_STDOUT.write(line + "\n")
    _REAL_STDOUT.flush()


def send_error(request_id: str | None, message: str):
    send_message({
        "type": "error",
        "id": request_id,
        "message": message,
    })


def _ast_contains_call_or_assignment(node: ast.AST) -> bool:
    """检查 AST 节点是否包含函数调用或赋值表达式（会产生副作用）。"""
    if isinstance(node, (ast.Call, ast.NamedExpr)):
        return True
    for child in ast.iter_child_nodes(node):
        if _ast_contains_call_or_assignment(child):
            return True
    return False


def _snapshot_workspace_files(root: str) -> dict[str, tuple[int, int]]:
    """记录 workspace 内文件的相对路径、修改时间和大小。"""
    snapshot = {}
    root_path = Path(root)
    if not root_path.exists():
        return snapshot

    skipped_dirs = {".git", "__pycache__"}
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in skipped_dirs]
        for name in filenames:
            if name.endswith((".pyc", ".pyo")):
                continue
            full_path = Path(dirpath) / name
            try:
                stat = full_path.stat()
                rel = full_path.relative_to(root_path).as_posix()
                snapshot[rel] = (stat.st_mtime_ns, stat.st_size)
            except Exception:
                continue

    return snapshot


def install_package(package_name: str) -> tuple[bool, str]:
    """使用 pip 安装包，返回 (success, output)。"""
    try:
        result = sp.run(
            [sys.executable, "-m", "pip", "install", package_name],
            capture_output=True,
            text=True,
            timeout=60,
        )
        output = result.stdout + result.stderr
        return result.returncode == 0, output
    except Exception as e:
        return False, str(e)


# ============================================================
# 文档解析 — 支持 .docx / .pdf / .xlsx / .txt / .csv / .md 等
# ============================================================

# 库名 → pip 包名 映射（有些库导入名和 pip 包名不一致）
_LIB_MAP = {
    "docx": "python-docx",
    "PyPDF2": "PyPDF2",
    "openpyxl": "openpyxl",
    "chardet": "chardet",
}


def _ensure_library(import_name: str):
    """确保 Python 库已安装，没有则自动 pip install。"""
    pip_name = _LIB_MAP.get(import_name, import_name)
    try:
        __import__(import_name)
    except ImportError:
        success, output = install_package(pip_name)
        if not success:
            raise ImportError(f"无法安装 {pip_name}:\n{output}")
        # 重新导入
        __import__(import_name)


def parse_document_file(file_path: str) -> str:
    """解析文档文件，返回提取的文本内容。

    支持格式：
      .txt / .md / .csv / .json / .html / .xml / .py / .js / .css 等 — 直接文本读取（自动检测编码）
      .docx  — Word 文档（需 python-docx）
      .pdf   — PDF 文档（需 PyPDF2）
      .xlsx / .xls  — Excel 表格（需 openpyxl）
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    if not path.is_file():
        raise ValueError(f"路径不是文件: {file_path}")

    ext = path.suffix.lower()
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > 50:
        raise ValueError(f"文件过大 ({size_mb:.1f} MB)，超过 50MB 上限")

    # ---- 纯文本类 ----
    if ext in (".txt", ".md", ".csv", ".json", ".html", ".htm", ".xml",
               ".yaml", ".yml", ".log", ".ini", ".cfg", ".toml",
               ".py", ".js", ".ts", ".css", ".sh", ".bat", ".ps1"):
        return _read_text_file(path)

    # ---- Word 文档 ----
    if ext in (".docx",):
        return _parse_docx(path)

    # ---- PDF 文档 ----
    if ext in (".pdf",):
        return _parse_pdf(path)

    # ---- Excel 表格 ----
    if ext in (".xlsx", ".xls", ".xlsm"):
        return _parse_xlsx(path)

    raise ValueError(f"不支持的文件格式: {ext}")


def _read_text_file(path: Path) -> str:
    """读取文本文件，自动检测编码。"""
    raw = path.read_bytes()
    # 尝试 UTF-8
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    # 尝试 chardet 自动检测
    try:
        _ensure_library("chardet")
        import chardet
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        return raw.decode(encoding, errors="replace")
    except Exception:
        pass
    # 最后尝试 GBK（Windows 中文常用）
    try:
        return raw.decode("gbk", errors="replace")
    except Exception:
        return raw.decode("utf-8", errors="replace")


def _parse_docx(path: Path) -> str:
    """解析 Word .docx 文档。"""
    _ensure_library("docx")
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 也尝试提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n\n".join(parts) if parts else "(文档无文本内容)"


def _parse_pdf(path: Path) -> str:
    """解析 PDF 文档。"""
    _ensure_library("PyPDF2")
    from PyPDF2 import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            parts.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
    return "\n\n".join(parts) if parts else "(PDF 无文本内容，可能是扫描件)"


def _parse_xlsx(path: Path) -> str:
    """解析 Excel 表格。"""
    _ensure_library("openpyxl")
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== 工作表: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(
                str(cell) if cell is not None else "" for cell in row
            )
            if row_text.strip():
                parts.append(row_text)
        if len(parts) > 5000:  # 大表格截断
            parts = parts[:5000]
            parts.append("...(表格内容过长，已截断)")
            break
    wb.close()
    return "\n".join(parts) if parts else "(表格无内容)"


class CodeExecutor:
    """在受限命名空间中执行 Python 代码。"""

    def __init__(self):
        self.globals = {
            "__builtins__": __builtins__,
            "__name__": "__exec__",
        }

    def execute(self, code: str, request_id: str, cwd: str | None = None):
        """
        执行 Python 代码，逐行输出通过 stream_output 实时发送。
        支持表达式结果返回。
        """
        original_cwd = os.getcwd()
        if cwd and os.path.isdir(cwd):
            os.chdir(cwd)
        before_files = _snapshot_workspace_files(cwd) if cwd and os.path.isdir(cwd) else {}

        # 捕获 stdout/stderr 实现实时流式输出
        old_stdout = sys.stdout
        old_stderr = sys.stderr
        captured_stdout = io.StringIO()
        captured_stderr = io.StringIO()

        class StreamProxy:
            """捕获输出并通过 JSON 协议实时发送，同时写入捕获流用于结果汇总。

            关键：不直接写入 sys.__stdout__，因为那会绕过 JSON-line 协议，
            导致原始文本与 JSON 消息混在一起。所有输出通过 send_message() 走协议通道。
            """
            def __init__(self, capture, stream_name, req_id):
                self.capture = capture
                self.stream_name = stream_name
                self.req_id = req_id
                self._buffer = ""

            def write(self, s):
                self.capture.write(s)
                # 逐行发送 stream_output
                self._buffer += s
                while "\n" in self._buffer:
                    line, self._buffer = self._buffer.split("\n", 1)
                    send_message({
                        "type": "stream_output",
                        "id": self.req_id,
                        "stream": self.stream_name,
                        "text": line + "\n",
                    })

            def flush(self):
                if self._buffer:
                    send_message({
                        "type": "stream_output",
                        "id": self.req_id,
                        "stream": self.stream_name,
                        "text": self._buffer,
                    })
                    self._buffer = ""

        sys.stdout = StreamProxy(captured_stdout, "stdout", request_id)
        sys.stderr = StreamProxy(captured_stderr, "stderr", request_id)

        result = None
        error = None
        files = []

        try:
            # 编译代码，尝试获取最后一个表达式的结果
            compiled = compile(code, "<execute>", "exec")
            exec(compiled, self.globals)

            # 尝试获取表达式结果：将最后一行作为表达式求值。
            # 但跳过会产生副作用的表达式（如函数调用 print()、赋值等），
            # 避免重复执行 exec 已经运行过的语句。
            last_line = code.strip().split("\n")[-1].strip()
            if last_line and not last_line.startswith(("def ", "class ", "if ", "for ", "while ", "try:", "with ", "import ", "from ", "elif ", "else:", "except", "finally:")):
                try:
                    expr_ast = ast.parse(last_line, mode="eval")
                    # 跳过函数调用（有副作用）和赋值表达式
                    if not _ast_contains_call_or_assignment(expr_ast):
                        expr_compiled = compile(expr_ast, "<expr>", "eval")
                        result = eval(expr_compiled, self.globals)
                except SyntaxError:
                    pass  # 最后一行不是表达式，忽略

        except Exception:
            error = traceback.format_exc()

        finally:
            sys.stdout.flush()
            sys.stderr.flush()
            sys.stdout = old_stdout
            sys.stderr = old_stderr

            # 收集执行期间生成的文件（相对路径）
            if cwd:
                try:
                    after_files = _snapshot_workspace_files(cwd)
                    for rel_path, meta in sorted(after_files.items()):
                        if rel_path not in before_files or before_files[rel_path] != meta:
                            files.append(rel_path)
                except Exception:
                    pass

            os.chdir(original_cwd)

        # 发送结果
        result_msg = {"type": "executeResult", "id": request_id}
        if error:
            result_msg["error"] = error
        else:
            result_msg["result"] = repr(result) if result is not None else None
            result_msg["stdout"] = captured_stdout.getvalue()
            result_msg["stderr"] = captured_stderr.getvalue()
            result_msg["files"] = files

        send_message(result_msg)


def main():
    # 发送就绪信号
    send_message({"type": "ready"})
    executor = CodeExecutor()
    cwd = os.getcwd()

    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue

        try:
            msg = json.loads(line)
        except json.JSONDecodeError as e:
            send_error(None, f"JSON 解析失败: {e}")
            continue

        msg_type = msg.get("type", "")
        request_id = msg.get("id", "")

        if msg_type == "shutdown":
            send_message({"type": "shutdownAck"})
            sys.exit(0)

        elif msg_type == "execute":
            code = msg.get("code", "")
            if not code:
                send_error(request_id, "execute 消息缺少 code 字段")
                continue

            # 支持指定工作目录
            work_dir = msg.get("cwd", cwd)
            executor.execute(code, request_id, work_dir)

        elif msg_type == "install":
            package = msg.get("package", "")
            if not package:
                send_error(request_id, "install 消息缺少 package 字段")
                continue

            success, output = install_package(package)
            send_message({
                "type": "executeResult",
                "id": request_id,
                "result": f"pip install {package}: {'成功' if success else '失败'}",
                "stdout": output,
            })

        elif msg_type == "parse_document":
            file_path = msg.get("filePath", "")
            if not file_path:
                send_error(request_id, "parse_document 消息缺少 filePath 字段")
                continue

            try:
                text = parse_document_file(file_path)
                send_message({
                    "type": "executeResult",
                    "id": request_id,
                    "result": text,
                    "stdout": "",
                    "stderr": "",
                    "files": [],
                })
            except Exception as e:
                send_message({
                    "type": "executeResult",
                    "id": request_id,
                    "error": f"文档解析失败: {e}",
                })

        elif msg_type == "cancel":
            # MVP: cancel 由 Node.js 侧 kill 进程实现，
            # 这里只是协议占位，实际取消不在这里处理。
            send_message({
                "type": "stream_output",
                "id": request_id,
                "stream": "stderr",
                "text": "[cancel] 取消请求由主进程处理\n",
            })

        else:
            send_error(request_id, f"未知消息类型: {msg_type}")


if __name__ == "__main__":
    main()
